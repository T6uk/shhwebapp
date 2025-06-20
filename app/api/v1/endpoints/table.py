# app/api/v1/endpoints/table.py
import csv
import datetime
import json
import logging
import mimetypes
import os
import platform
import re
import subprocess
import time
import traceback
import uuid
from datetime import date
from datetime import datetime
from io import StringIO
from typing import Optional

import orjson
from fastapi import APIRouter, HTTPException, Depends, Request
from fastapi import Query, Response, Form
from sqlalchemy import or_
from sqlalchemy import text
from sqlalchemy.ext.asyncio import AsyncSession
from sqlalchemy.orm import Session  # Add this import

from app.api.dependencies import get_current_active_user
from app.core.cache import get_cache, set_cache
from app.core.db import get_db
from app.core.user_db import get_user_db
from app.models.saved_filter import SavedFilter
from app.models.table import BigTable
from app.models.user import User
from app.services.edit_service import (
    verify_edit_permission, get_editable_columns, update_cell_value,
    get_session_changes, undo_change, check_for_changes
)

# Make sure these are imported for the Koondaja functionality

# Set up logging
logging.basicConfig(level=logging.INFO)
logger = logging.getLogger(__name__)

router = APIRouter(prefix="/table", tags=["table"])


@router.get("/data")
async def get_table_data(
        request: Request,
        start_row: int = Query(0, ge=0),
        end_row: int = Query(100, ge=1),
        search: Optional[str] = None,
        sort_field: Optional[str] = None,
        sort_dir: Optional[str] = None,
        filter_model: Optional[str] = None,
        db: AsyncSession = Depends(get_db)
):
    """
    Get paginated table data with filtering and sorting
    """
    start_time = time.time()
    try:
        # Detect database type
        from app.core.db import is_using_local_db
        using_sqlite = is_using_local_db()

        # Build SQL query components
        where_clauses = []
        query_params = {}

        # Get column data types for proper type conversion - cache this to improve performance
        column_types = {}
        try:
            if using_sqlite:
                # Get column types from SQLite
                schema_query = "PRAGMA table_info('taitur_data')"
                schema_result = await db.execute(text(schema_query))

                for row in schema_result:
                    column_name = row[1]  # name is at index 1
                    data_type = row[2]  # type is at index 2
                    column_types[column_name] = {
                        "data_type": data_type.lower(),
                        "udt_name": data_type.lower()
                    }
            else:
                # Get table column information including data types from PostgreSQL
                schema_query = """
                    SELECT column_name, data_type, udt_name 
                    FROM information_schema.columns 
                    WHERE table_name = :table_name
                """
                schema_result = await db.execute(text(schema_query), {"table_name": BigTable.name})

                for row in schema_result:
                    column_name = row[0]
                    data_type = row[1]
                    udt_name = row[2]
                    column_types[column_name] = {
                        "data_type": data_type,
                        "udt_name": udt_name
                    }

            logger.info(f"Retrieved column types: {column_types}")
        except Exception as e:
            logger.error(f"Error getting column types: {str(e)}")
            # Continue without type information - we'll try to guess types

        # Log received filter model for debugging
        logger.info(f"Received filter_model: {filter_model}")

        # Add search filters if provided
        if filter_model:
            try:
                # Parse the filter model
                try:
                    filters = json.loads(filter_model)
                    logger.info(f"Successfully parsed filter_model: {filters}")
                except json.JSONDecodeError:
                    logger.error(f"Invalid JSON in filter_model: {filter_model}")
                    raise ValueError("Invalid filter model format")

                # Process each filter
                filter_idx = 0
                for field, filter_config in filters.items():
                    filter_type = filter_config.get('type')
                    filter_value = filter_config.get('filter')

                    logger.info(f"Processing filter: field={field}, type={filter_type}, value={filter_value}")

                    # Skip if field doesn't exist in the table
                    if field not in column_types and field != "id":  # Always allow "id"
                        logger.warning(f"Field {field} not found in table schema, skipping filter")
                        continue

                    # Determine the appropriate cast expression based on DB type
                    cast_expr = "" if using_sqlite else "::text"

                    try:
                        # Handle each filter type with better type handling
                        if filter_type == 'contains':
                            param_name = f"filter_{filter_idx}"
                            where_clauses.append(f'"{field}"{cast_expr} LIKE :{param_name}')
                            query_params[param_name] = f"%{filter_value}%"
                            filter_idx += 1

                        elif filter_type == 'notContains':
                            param_name = f"filter_{filter_idx}"
                            where_clauses.append(f'("{field}" IS NULL OR "{field}"{cast_expr} NOT LIKE :{param_name})')
                            query_params[param_name] = f"%{filter_value}%"
                            filter_idx += 1

                        elif filter_type == 'equals':
                            param_name = f"filter_{filter_idx}"
                            where_clauses.append(f'"{field}" = :{param_name}')

                            # Try to convert values based on column type
                            col_type_info = column_types.get(field, {})
                            col_type = col_type_info.get("data_type", "").lower()

                            # Handle numeric types
                            if field == "id" or col_type in ("integer", "int", "bigint", "smallint"):
                                try:
                                    query_params[param_name] = int(filter_value)
                                except (ValueError, TypeError):
                                    query_params[param_name] = filter_value
                            elif col_type in ("numeric", "decimal", "real", "double", "float"):
                                try:
                                    query_params[param_name] = float(filter_value)
                                except (ValueError, TypeError):
                                    query_params[param_name] = filter_value
                            # Handle date types
                            elif col_type in ("date", "timestamp"):
                                query_params[param_name] = filter_value
                            else:
                                query_params[param_name] = filter_value

                            filter_idx += 1

                        elif filter_type == 'notEqual':
                            param_name = f"filter_{filter_idx}"
                            where_clauses.append(f'("{field}" IS NULL OR "{field}" != :{param_name})')

                            # Type conversion similar to equals
                            col_type_info = column_types.get(field, {})
                            col_type = col_type_info.get("data_type", "").lower()

                            if field == "id" or col_type in ("integer", "int", "bigint", "smallint"):
                                try:
                                    query_params[param_name] = int(filter_value)
                                except (ValueError, TypeError):
                                    query_params[param_name] = filter_value
                            elif col_type in ("numeric", "decimal", "real", "double", "float"):
                                try:
                                    query_params[param_name] = float(filter_value)
                                except (ValueError, TypeError):
                                    query_params[param_name] = filter_value
                            else:
                                query_params[param_name] = filter_value

                            filter_idx += 1

                        elif filter_type == 'startsWith':
                            param_name = f"filter_{filter_idx}"
                            where_clauses.append(f'"{field}"{cast_expr} LIKE :{param_name}')
                            query_params[param_name] = f"{filter_value}%"
                            filter_idx += 1

                        elif filter_type == 'endsWith':
                            param_name = f"filter_{filter_idx}"
                            where_clauses.append(f'"{field}"{cast_expr} LIKE :{param_name}')
                            query_params[param_name] = f"%{filter_value}"
                            filter_idx += 1

                        elif filter_type == 'notBlank':
                            # Improved notBlank filter to correctly filter out ALL types of blank values
                            if using_sqlite:
                                # SQLite implementation - more robust check including whitespace-only strings
                                where_clauses.append(
                                    f'("{field}" IS NOT NULL AND "{field}" != "" AND TRIM("{field}") != "")')
                                logger.info(f"Added improved notBlank filter for {field} using SQLite syntax")
                            else:
                                # PostgreSQL implementation - handle all types of blank values and type-specific checks
                                col_type_info = column_types.get(field, {})
                                col_type = col_type_info.get("data_type", "").lower()
                                if col_type in (
                                        "integer", "int", "bigint", "smallint", "numeric", "decimal", "real", "double",
                                        "float"):
                                    # For numeric types, IS NOT NULL is sufficient
                                    where_clauses.append(f'"{field}" IS NOT NULL')
                                elif col_type in ("date", "timestamp", "timestamp with time zone"):
                                    # For date/timestamp types, IS NOT NULL is sufficient
                                    where_clauses.append(f'"{field}" IS NOT NULL')
                                else:
                                    # For string and other types, check NULL, empty string, and whitespace-only
                                    where_clauses.append(
                                        f'("{field}" IS NOT NULL AND "{field}"::text != \'\' AND TRIM(COALESCE("{field}"::text, \'\')) != \'\')')
                                logger.info(
                                    f"Added improved notBlank filter for {field} using PostgreSQL syntax for type {col_type}")

                        elif filter_type == 'blank':
                            # Improved blank filter to correctly match ALL types of blank values
                            if using_sqlite:
                                # SQLite implementation - including whitespace-only strings
                                where_clauses.append(f'("{field}" IS NULL OR "{field}" = "" OR TRIM("{field}") = "")')
                                logger.info(f"Added improved blank filter for {field} using SQLite syntax")
                            else:
                                # PostgreSQL implementation - type-specific handling
                                col_type_info = column_types.get(field, {})
                                col_type = col_type_info.get("data_type", "").lower()

                                if col_type in (
                                        "integer", "int", "bigint", "smallint", "numeric", "decimal", "real", "double",
                                        "float"):
                                    # For numeric types, IS NULL is sufficient
                                    where_clauses.append(f'"{field}" IS NULL')
                                elif col_type in ("date", "timestamp", "timestamp with time zone"):
                                    # For date/timestamp types, IS NULL is sufficient
                                    where_clauses.append(f'"{field}" IS NULL')
                                else:
                                    # For string and other types, check NULL, empty string, and whitespace-only
                                    where_clauses.append(
                                        f'("{field}" IS NULL OR "{field}"::text = \'\' OR TRIM(COALESCE("{field}"::text, \'\')) = \'\')')

                                logger.info(
                                    f"Added improved blank filter for {field} using PostgreSQL syntax for type {col_type}")

                        elif filter_type == 'greaterThan':
                            param_name = f"filter_{filter_idx}"
                            where_clauses.append(f'"{field}" > :{param_name}')

                            # Type conversion
                            col_type_info = column_types.get(field, {})
                            col_type = col_type_info.get("data_type", "").lower()

                            if col_type in ("integer", "int", "bigint", "smallint"):
                                try:
                                    query_params[param_name] = int(filter_value)
                                except (ValueError, TypeError):
                                    query_params[param_name] = filter_value
                            elif col_type in ("numeric", "decimal", "real", "double", "float"):
                                try:
                                    query_params[param_name] = float(filter_value)
                                except (ValueError, TypeError):
                                    query_params[param_name] = filter_value
                            else:
                                query_params[param_name] = filter_value

                            filter_idx += 1

                        elif filter_type == 'greaterThanOrEqual':
                            param_name = f"filter_{filter_idx}"
                            where_clauses.append(f'"{field}" >= :{param_name}')

                            # Type conversion
                            col_type_info = column_types.get(field, {})
                            col_type = col_type_info.get("data_type", "").lower()

                            if col_type in ("integer", "int", "bigint", "smallint"):
                                try:
                                    query_params[param_name] = int(filter_value)
                                except (ValueError, TypeError):
                                    query_params[param_name] = filter_value
                            elif col_type in ("numeric", "decimal", "real", "double", "float"):
                                try:
                                    query_params[param_name] = float(filter_value)
                                except (ValueError, TypeError):
                                    query_params[param_name] = filter_value
                            else:
                                query_params[param_name] = filter_value

                            filter_idx += 1

                        elif filter_type == 'lessThan':
                            param_name = f"filter_{filter_idx}"
                            where_clauses.append(f'"{field}" < :{param_name}')

                            # Type conversion
                            col_type_info = column_types.get(field, {})
                            col_type = col_type_info.get("data_type", "").lower()

                            if col_type in ("integer", "int", "bigint", "smallint"):
                                try:
                                    query_params[param_name] = int(filter_value)
                                except (ValueError, TypeError):
                                    query_params[param_name] = filter_value
                            elif col_type in ("numeric", "decimal", "real", "double", "float"):
                                try:
                                    query_params[param_name] = float(filter_value)
                                except (ValueError, TypeError):
                                    query_params[param_name] = filter_value
                            else:
                                query_params[param_name] = filter_value

                            filter_idx += 1

                        elif filter_type == 'lessThanOrEqual':
                            param_name = f"filter_{filter_idx}"
                            where_clauses.append(f'"{field}" <= :{param_name}')

                            # Type conversion
                            col_type_info = column_types.get(field, {})
                            col_type = col_type_info.get("data_type", "").lower()

                            if col_type in ("integer", "int", "bigint", "smallint"):
                                try:
                                    query_params[param_name] = int(filter_value)
                                except (ValueError, TypeError):
                                    query_params[param_name] = filter_value
                            elif col_type in ("numeric", "decimal", "real", "double", "float"):
                                try:
                                    query_params[param_name] = float(filter_value)
                                except (ValueError, TypeError):
                                    query_params[param_name] = filter_value
                            else:
                                query_params[param_name] = filter_value

                            filter_idx += 1

                        elif filter_type == 'inRange' and isinstance(filter_value, dict):
                            from_value = filter_value.get('from')
                            to_value = filter_value.get('to')

                            # Type conversion
                            col_type_info = column_types.get(field, {})
                            col_type = col_type_info.get("data_type", "").lower()

                            # Only add range conditions if values are provided
                            if from_value is not None and from_value != "":
                                param_name = f"filter_{filter_idx}"
                                where_clauses.append(f'"{field}" >= :{param_name}')

                                if col_type in ("integer", "int", "bigint", "smallint"):
                                    try:
                                        query_params[param_name] = int(from_value)
                                    except (ValueError, TypeError):
                                        query_params[param_name] = from_value
                                elif col_type in ("numeric", "decimal", "real", "double", "float"):
                                    try:
                                        query_params[param_name] = float(from_value)
                                    except (ValueError, TypeError):
                                        query_params[param_name] = from_value
                                elif col_type in ("date", "timestamp"):
                                    query_params[param_name] = from_value
                                else:
                                    query_params[param_name] = from_value

                                filter_idx += 1

                            if to_value is not None and to_value != "":
                                param_name = f"filter_{filter_idx}"
                                where_clauses.append(f'"{field}" <= :{param_name}')

                                if col_type in ("integer", "int", "bigint", "smallint"):
                                    try:
                                        query_params[param_name] = int(to_value)
                                    except (ValueError, TypeError):
                                        query_params[param_name] = to_value
                                elif col_type in ("numeric", "decimal", "real", "double", "float"):
                                    try:
                                        query_params[param_name] = float(to_value)
                                    except (ValueError, TypeError):
                                        query_params[param_name] = to_value
                                elif col_type in ("date", "timestamp"):
                                    query_params[param_name] = to_value
                                else:
                                    query_params[param_name] = to_value

                                filter_idx += 1

                        else:
                            logger.warning(f"Unsupported filter type: {filter_type}")

                        logger.info(f"Successfully processed filter: {field} {filter_type}")

                    except Exception as e:
                        logger.error(f"Error processing specific filter '{field}' with type '{filter_type}': {str(e)}")
                        logger.error(traceback.format_exc())
                        # Continue with other filters but log the error

            except Exception as e:
                logger.error(f"Error processing filter model: {str(e)}", exc_info=True)
                # Continue without failing - best effort filtering

            # After processing filters, log complete SQL info
            if where_clauses:
                logger.info(f"Final WHERE clauses: {where_clauses}")
                logger.info(f"Final query params: {query_params}")

        # Combine WHERE clauses
        where_sql = ""
        if where_clauses:
            where_sql = f" WHERE {' AND '.join(where_clauses)}"

        # Log the final SQL for debugging
        logger.info(f"WHERE clause: {where_sql}")
        logger.info(f"Query params: {query_params}")

        # Build and execute count query
        count_sql = f'SELECT COUNT(*) FROM "{BigTable.name}"{where_sql}'
        count_result = await db.execute(text(count_sql), query_params)
        total_rows = count_result.scalar() or 0

        # Build and execute data query with sorting and pagination
        data_sql = f'SELECT * FROM "{BigTable.name}"{where_sql}'
        if sort_field:
            sort_direction = "DESC" if sort_dir and sort_dir.lower() == "desc" else "ASC"
            data_sql += f' ORDER BY "{sort_field}" {sort_direction}'
        else:
            data_sql += ' ORDER BY 1'  # Default sort

        # Add pagination
        data_sql += ' LIMIT :limit OFFSET :offset'
        query_params["limit"] = end_row - start_row
        query_params["offset"] = start_row

        # Execute data query
        result = await db.execute(text(data_sql), query_params)
        rows = result.fetchall()

        # Convert to list of dicts for JSON response
        data = []
        for row in rows:
            row_dict = {}
            for key in row._mapping.keys():
                value = row._mapping[key]
                # Handle datetime objects for JSON serialization
                if isinstance(value, (datetime, date)):
                    row_dict[str(key)] = value.isoformat()
                else:
                    row_dict[str(key)] = value
            data.append(row_dict)

        logger.info(f"Query returned {len(data)} rows out of {total_rows} total in {time.time() - start_time:.3f}s")

        # Create response with row count info
        response_data = {
            "rowData": data,
            "rowCount": total_rows,
            "startRow": start_row,
            "endRow": min(start_row + len(data), total_rows)
        }

        # Return using orjson for faster serialization
        return Response(
            content=orjson.dumps(response_data),
            media_type="application/json"
        )

    except Exception as e:
        logger.exception(f"Error getting table data: {str(e)}")
        raise HTTPException(status_code=500, detail=f"Database error: {str(e)}")


@router.get("/columns")
async def get_columns(
        db: AsyncSession = Depends(get_db),
        response: Response = None
):
    """
    Get table columns metadata for frontend with support for both PostgreSQL and SQLite
    """
    start_time = time.time()
    try:
        # Try to get from cache
        cached_columns = await get_cache("table_columns")
        if cached_columns:
            logger.info(f"Returning cached columns in {time.time() - start_time:.3f}s")
            # Set cache control header for browser caching
            if response:
                response.headers["Cache-Control"] = "public, max-age=86400"
            return Response(
                content=orjson.dumps(cached_columns),
                media_type="application/json",
                headers={"Cache-Control": "public, max-age=86400"}
            )

        # Detect database type by checking connection driver
        from app.core.db import is_using_local_db
        using_sqlite = is_using_local_db()

        columns = []

        if using_sqlite:
            # SQLite schema query using PRAGMA
            sql = "PRAGMA table_info('taitur_data')"
            result = await db.execute(text(sql))
            rows = result.fetchall()

            # Process SQLite column info
            # PRAGMA table_info returns: cid, name, type, notnull, dflt_value, pk
            for row in rows:
                column_info = {
                    "field": row[1],  # name
                    "title": row[1].capitalize().replace("_", " "),
                    "type": row[2],  # type
                    "nullable": row[3] == 0,  # notnull (0 = nullable, 1 = not null)
                    "hasDefault": row[4] is not None,  # dflt_value
                    "primary_key": row[5] == 1  # pk
                }
                columns.append(column_info)
        else:
            # PostgreSQL schema query using information_schema
            sql = f"""
                SELECT 
                    column_name,
                    data_type,
                    column_default,
                    is_nullable
                FROM 
                    information_schema.columns
                WHERE 
                    table_name = '{BigTable.name}'
                ORDER BY 
                    ordinal_position
            """

            result = await db.execute(text(sql))
            rows = result.fetchall()

            # Process PostgreSQL column info
            for row in rows:
                column_name = row[0]
                data_type = row[1]
                column_default = row[2]
                is_nullable = row[3]

                column_info = {
                    "field": column_name,
                    "title": column_name.capitalize().replace("_", " "),
                    "type": data_type,
                    "nullable": is_nullable == "YES",
                    "hasDefault": column_default is not None
                }
                columns.append(column_info)

        logger.info(f"Returning {len(columns)} columns in {time.time() - start_time:.3f}s")

        # Create response and cache it (24 hour TTL for schema info)
        response_data = {"columns": columns}
        await set_cache("table_columns", response_data, expire=86400)

        # Set cache control header for browser caching
        return Response(
            content=orjson.dumps(response_data),
            media_type="application/json",
            headers={"Cache-Control": "public, max-age=86400"}
        )

    except Exception as e:
        logger.exception(f"Error getting columns: {str(e)}")
        raise HTTPException(status_code=500, detail=f"Database error: {str(e)}")


@router.get("/editable-columns")
async def get_editable_columns_endpoint(
        db: Session = Depends(get_user_db),
        current_user: User = Depends(get_current_active_user)
):
    """Get which columns are editable for the current user"""
    editable_columns = await get_editable_columns(db, current_user)

    return {
        "columns": editable_columns,
        "can_edit": current_user.can_edit
    }


@router.post("/verify-edit-password")
async def verify_edit_password_endpoint(
        password: str = Form(...),
        current_user: User = Depends(get_current_active_user)
):
    """Verify the edit password to enable editing mode"""
    is_valid = await verify_edit_permission(current_user, password)

    if is_valid:
        # Generate a session ID for tracking changes
        session_id = str(uuid.uuid4())
        return {
            "success": True,
            "session_id": session_id,
            "message": "Edit mode enabled successfully"
        }

    return {
        "success": False,
        "message": "Invalid password or you don't have edit permission"
    }


@router.post("/update-cell")
async def update_cell_endpoint(
        request: Request,
        table_name: str = Form(...),
        row_id: str = Form(...),
        column_name: str = Form(...),
        old_value: Optional[str] = Form(None),
        new_value: str = Form(...),
        session_id: str = Form(...),
        db: AsyncSession = Depends(get_db),
        user_db: Session = Depends(get_user_db),
        current_user: User = Depends(get_current_active_user)
):
    """Update a cell value with change tracking"""
    success = await update_cell_value(
        db,
        user_db,
        current_user,
        table_name,
        row_id,
        column_name,
        old_value,
        new_value,
        session_id,
        client_ip=request.client.host,
        user_agent=request.headers.get("User-Agent", "")
    )

    return {"success": success}


@router.get("/session-changes/{session_id}")
async def get_session_changes_endpoint(
        session_id: str,
        user_db: Session = Depends(get_user_db),
        current_user: User = Depends(get_current_active_user)
):
    """Get all changes made in a session for undo functionality"""
    changes = await get_session_changes(user_db, current_user, session_id)

    return {
        "changes": [
            {
                "id": change.id,
                "table_name": change.table_name,
                "row_id": change.row_id,
                "column_name": change.column_name,
                "old_value": change.old_value,
                "new_value": change.new_value,
                "changed_at": change.changed_at.isoformat()
            }
            for change in changes
        ]
    }


@router.post("/undo-change/{change_id}")
async def undo_change_endpoint(
        request: Request,
        change_id: int,
        db: AsyncSession = Depends(get_db),
        user_db: Session = Depends(get_user_db),
        current_user: User = Depends(get_current_active_user)
):
    """Undo a specific change"""
    success = await undo_change(
        db,
        user_db,
        current_user,
        change_id,
        client_ip=request.client.host,
        user_agent=request.headers.get("User-Agent", "")
    )

    return {"success": success}


@router.get("/check-for-changes")
async def check_for_changes_endpoint(
        last_checked: Optional[str] = Query(None),
        user_db: Session = Depends(get_user_db),
        current_user: User = Depends(get_current_active_user)
):
    """Check if there are changes from other users since last check"""
    last_checked_time = None
    if last_checked:
        try:
            last_checked_time = datetime.fromisoformat(last_checked)
        except ValueError:
            pass

    result = await check_for_changes(user_db, current_user, last_checked_time)
    return result


@router.get("/filter-values/{column_name}")
async def get_filter_values(
        column_name: str,
        search: Optional[str] = None,
        limit: int = Query(100, ge=1, le=1000),
        db: AsyncSession = Depends(get_db)
):
    """Get available values for a column to populate filter dropdowns"""
    try:
        from app.services.data_loader import get_available_filter_values

        values = await get_available_filter_values(db, column_name, search, limit)

        return {
            "values": values,
            "count": len(values)
        }

    except Exception as e:
        logger.exception(f"Error getting filter values: {str(e)}")
        raise HTTPException(status_code=500, detail=f"Error getting filter values: {str(e)}")


@router.post("/save-filter")
async def save_filter(
        name: str = Form(...),
        description: Optional[str] = Form(None),
        filter_model: str = Form(...),
        is_public: bool = Form(False),
        db: Session = Depends(get_user_db),
        current_user: User = Depends(get_current_active_user)
):
    """Save a filter configuration for later use"""
    try:
        # Create a new filter model
        new_filter = SavedFilter(
            name=name,
            description=description,
            filter_model=filter_model,
            is_public=is_public,
            user_id=current_user.id,
            created_at=datetime.utcnow()
        )

        db.add(new_filter)
        db.commit()
        db.refresh(new_filter)

        return {
            "id": new_filter.id,
            "name": new_filter.name,
            "message": "Filter saved successfully"
        }

    except Exception as e:
        logger.exception(f"Error saving filter: {str(e)}")
        raise HTTPException(status_code=500, detail=f"Error saving filter: {str(e)}")


@router.get("/saved-filters")
async def get_saved_filters(
        db: Session = Depends(get_user_db),
        current_user: User = Depends(get_current_active_user)
):
    """Get all saved filters for the current user plus public filters"""
    try:
        # Get user's filters and public filters
        filters = db.query(SavedFilter).filter(
            or_(
                SavedFilter.user_id == current_user.id,
                SavedFilter.is_public == True
            )
        ).order_by(SavedFilter.name).all()

        return {
            "filters": [
                {
                    "id": f.id,
                    "name": f.name,
                    "description": f.description,
                    "is_public": f.is_public,
                    "created_at": f.created_at.isoformat(),
                    "is_owner": f.user_id == current_user.id
                }
                for f in filters
            ]
        }

    except Exception as e:
        logger.exception(f"Error getting saved filters: {str(e)}")
        raise HTTPException(status_code=500, detail=f"Error getting saved filters: {str(e)}")


@router.get("/saved-filter/{filter_id}")
async def get_saved_filter(
        filter_id: int,
        db: Session = Depends(get_user_db),
        current_user: User = Depends(get_current_active_user)
):
    """Get a specific saved filter"""
    try:
        # Get the filter
        filter = db.query(SavedFilter).filter(SavedFilter.id == filter_id).first()

        if not filter:
            raise HTTPException(status_code=404, detail="Filter not found")

        # Check if user has access (owner or public filter)
        if filter.user_id != current_user.id and not filter.is_public:
            raise HTTPException(status_code=403, detail="Access denied to this filter")

        return {
            "id": filter.id,
            "name": filter.name,
            "description": filter.description,
            "filter_model": filter.filter_model,
            "is_public": filter.is_public,
            "created_at": filter.created_at.isoformat(),
            "is_owner": filter.user_id == current_user.id
        }

    except HTTPException:
        raise
    except Exception as e:
        logger.exception(f"Error getting filter details: {str(e)}")
        raise HTTPException(status_code=500, detail=f"Error getting filter details: {str(e)}")


@router.delete("/saved-filter/{filter_id}")
async def delete_saved_filter(
        filter_id: int,
        db: Session = Depends(get_user_db),
        current_user: User = Depends(get_current_active_user)
):
    """Delete a saved filter"""
    try:
        # Get the filter
        filter = db.query(SavedFilter).filter(SavedFilter.id == filter_id).first()

        if not filter:
            raise HTTPException(status_code=404, detail="Filter not found")

        # Check if user is the owner
        if filter.user_id != current_user.id:
            raise HTTPException(status_code=403, detail="You can only delete your own filters")

        # Delete the filter
        db.delete(filter)
        db.commit()

        return {
            "message": "Filter deleted successfully"
        }

    except HTTPException:
        raise
    except Exception as e:
        logger.exception(f"Error deleting filter: {str(e)}")
        raise HTTPException(status_code=500, detail=f"Error deleting filter: {str(e)}")


@router.get("/open-folder/{toimiku_nr}")
async def open_folder(
        toimiku_nr: str,
        current_user: User = Depends(get_current_active_user)
):
    """Open a folder for the specified toimiku_nr with improved path sanitization"""
    try:
        # Log input parameter
        logger.info(f"Received request to open folder for toimiku_nr: {toimiku_nr}")

        # Sanitize the toimiku_nr to ensure valid folder name
        # Replace any characters not allowed in Windows filenames
        sanitized_nr = re.sub(r'[\\/:*?"<>|]', '_', toimiku_nr)
        logger.info(f"Sanitized toimiku_nr: {sanitized_nr}")

        # Create the folder path ensuring correct Windows path format
        folder_path = os.path.join("c:\\", "virtuaaltoimik", toimiku_nr)
        logger.info(f"Target folder path: {folder_path}")

        # Check if folder exists
        if not os.path.exists(folder_path):
            logger.info(f"Folder does not exist, attempting to create: {folder_path}")
            try:
                os.makedirs(folder_path, exist_ok=True)
                logger.info(f"Successfully created folder: {folder_path}")
            except Exception as folder_error:
                error_msg = f"Failed to create folder: {str(folder_error)}"
                logger.error(error_msg)
                return {"success": False, "message": error_msg, "path": folder_path}
        else:
            logger.info(f"Folder already exists: {folder_path}")

        # Get system info for debugging
        system = platform.system()
        logger.info(f"Operating system: {system}")

        # Open the folder based on platform
        try:
            if system == "Windows":
                logger.info(f"Using os.startfile() to open folder on Windows")
                os.startfile(folder_path)
            elif system == "Darwin":  # macOS
                logger.info(f"Using 'open' command to open folder on macOS")
                result = subprocess.Popen(["open", folder_path], stdout=subprocess.PIPE, stderr=subprocess.PIPE)
                stdout, stderr = result.communicate()
                logger.info(f"Command output - stdout: {stdout}, stderr: {stderr}")
            elif system == "Linux":
                logger.info(f"Using 'xdg-open' command to open folder on Linux")
                result = subprocess.Popen(["xdg-open", folder_path], stdout=subprocess.PIPE, stderr=subprocess.PIPE)
                stdout, stderr = result.communicate()
                logger.info(f"Command output - stdout: {stdout}, stderr: {stderr}")
            else:
                error_msg = f"Unsupported platform: {system}"
                logger.error(error_msg)
                return {"success": False, "message": error_msg, "path": folder_path}

            logger.info(f"Successfully opened folder: {folder_path}")
            return {
                "success": True,
                "message": f"Folder opened: {folder_path}",
                "path": folder_path,
                "system": system
            }

        except Exception as open_error:
            error_detail = f"Error opening folder: {str(open_error)}"
            logger.error(error_detail)
            logger.error(traceback.format_exc())
            return {
                "success": False,
                "message": error_detail,
                "path": folder_path,
                "system": system,
                "error": str(open_error)
            }

    except Exception as e:
        error_trace = traceback.format_exc()
        logger.exception(f"Unexpected error in open_folder endpoint: {str(e)}")
        logger.error(error_trace)
        return {
            "success": False,
            "message": f"Unexpected error: {str(e)}",
            "trace": error_trace
        }


@router.get("/folder-contents/{toimiku_nr}")
async def get_folder_contents(
        toimiku_nr: str,
        current_user: User = Depends(get_current_active_user)
):
    """Get the contents of a specific toimiku folder"""
    try:
        # Sanitize the toimiku_nr to ensure valid folder name
        sanitized_nr = re.sub(r'[\\/:*?"<>|]', '_', toimiku_nr)
        logger.info(f"Getting contents for folder: {sanitized_nr}")

        # Create the folder path
        folder_path = os.path.join("c:\\", "virtuaaltoimik", sanitized_nr)

        # Create folder if it doesn't exist
        if not os.path.exists(folder_path):
            os.makedirs(folder_path, exist_ok=True)
            logger.info(f"Created folder: {folder_path}")

        # Get list of files and directories
        items = []

        try:
            for item in os.listdir(folder_path):
                item_path = os.path.join(folder_path, item)
                is_directory = os.path.isdir(item_path)

                # Get file stats
                stats = os.stat(item_path)

                # Guess the mime type if it's a file
                mime_type = None
                if not is_directory:
                    mime_type, _ = mimetypes.guess_type(item_path)

                # Format the modified time - using strftime instead of fromisoformat
                modified_time = datetime.fromtimestamp(stats.st_mtime).strftime("%d.%m.%Y %H:%M")

                # Add item info to the list
                items.append({
                    "name": item,
                    "is_directory": is_directory,
                    "size": stats.st_size if not is_directory else 0,
                    "formatted_size": format_file_size(stats.st_size) if not is_directory else "",
                    "modified": modified_time,
                    "extension": os.path.splitext(item)[1].lower() if not is_directory else "",
                    "mime_type": mime_type,
                    "path": item_path
                })

            # Sort items: directories first, then files alphabetically
            items.sort(key=lambda x: (not x["is_directory"], x["name"].lower()))

            return {
                "success": True,
                "path": folder_path,
                "toimiku_nr": sanitized_nr,
                "items": items
            }

        except Exception as e:
            logger.error(f"Error listing directory: {str(e)}")
            return {
                "success": False,
                "path": folder_path,
                "toimiku_nr": sanitized_nr,
                "error": str(e),
                "items": []
            }

    except Exception as e:
        logger.exception(f"Error in get_folder_contents: {str(e)}")
        raise HTTPException(status_code=500, detail=f"Error getting folder contents: {str(e)}")


def format_file_size(size_bytes):
    """Format file size in human readable format"""
    if size_bytes == 0:
        return "0 B"
    size_names = ["B", "KB", "MB", "GB", "TB"]
    import math
    i = int(math.floor(math.log(size_bytes, 1024)))
    p = math.pow(1024, i)
    s = round(size_bytes / p, 2)
    return f"{s} {size_names[i]}"


@router.post("/file-operation")
async def file_operation(
        operation: str = Form(...),
        path: str = Form(...),
        new_name: Optional[str] = Form(None),
        current_user: User = Depends(get_current_active_user)
):
    """Perform a file operation (delete, rename, etc.)"""
    try:
        logger.info(f"File operation: {operation} on path: {path}")

        # Verify the path is within the virtuaaltoimik directory for security
        if not path.startswith(r'c:\virtuaaltoimik'):
            return {
                "success": False,
                "message": "Keelatud operatsioon: vale kaust"
            }

        # Check if path exists
        if not os.path.exists(path):
            return {
                "success": False,
                "message": f"Faili või kausta ei leitud: {path}"
            }

        # Perform the requested operation
        if operation == "delete":
            if os.path.isdir(path):
                import shutil
                shutil.rmtree(path)
                logger.info(f"Deleted directory: {path}")
            else:
                os.remove(path)
                logger.info(f"Deleted file: {path}")

            return {
                "success": True,
                "message": "Fail või kaust kustutatud"
            }

        elif operation == "rename" and new_name:
            # Sanitize the new name
            sanitized_name = re.sub(r'[\\/:*?"<>|]', '_', new_name)

            # Create the new path
            directory = os.path.dirname(path)
            new_path = os.path.join(directory, sanitized_name)

            # Check if target already exists
            if os.path.exists(new_path):
                return {
                    "success": False,
                    "message": f"Fail või kaust '{sanitized_name}' on juba olemas"
                }

            # Rename the file or directory
            os.rename(path, new_path)
            logger.info(f"Renamed: {path} to {new_path}")

            return {
                "success": True,
                "message": "Fail või kaust ümbernimetatud",
                "new_path": new_path
            }

        else:
            return {
                "success": False,
                "message": f"Tundmatu operatsioon: {operation}"
            }

    except Exception as e:
        logger.exception(f"Error in file operation: {str(e)}")
        return {
            "success": False,
            "message": f"Viga failioperatsioonis: {str(e)}"
        }


@router.get("/document-templates")
async def get_document_templates(
        current_user: User = Depends(get_current_active_user)
):
    """Get document templates from the templates directory"""
    try:
        # Define templates directory
        templates_dir = r"C:\Taitemenetlus\uksikdokumendid\dokumendipohjad"
        logger.info(f"Getting document templates from: {templates_dir}")

        # Create directory if it doesn't exist
        if not os.path.exists(templates_dir):
            os.makedirs(templates_dir, exist_ok=True)
            logger.info(f"Created templates directory: {templates_dir}")

        # Get list of template files
        templates = []

        try:
            for item in os.listdir(templates_dir):
                item_path = os.path.join(templates_dir, item)

                # Skip directories, we only want files
                if os.path.isdir(item_path):
                    continue

                # Get file stats
                stats = os.stat(item_path)

                # Get file extension
                _, extension = os.path.splitext(item)

                # Format the modified time
                modified_time = datetime.fromtimestamp(stats.st_mtime).strftime("%d.%m.%Y %H:%M")

                # Add template info to the list
                templates.append({
                    "name": item,
                    "path": item_path,
                    "size": stats.st_size,
                    "formatted_size": format_file_size(stats.st_size),
                    "modified": modified_time,
                    "extension": extension.lower()
                })

            # Sort templates by name
            templates.sort(key=lambda x: x["name"].lower())

            return {
                "success": True,
                "templates_dir": templates_dir,
                "templates": templates
            }

        except Exception as e:
            logger.error(f"Error listing templates directory: {str(e)}")
            return {
                "success": False,
                "templates_dir": templates_dir,
                "error": str(e),
                "templates": []
            }

    except Exception as e:
        logger.exception(f"Error in get_document_templates: {str(e)}")
        raise HTTPException(status_code=500, detail=f"Error getting document templates: {str(e)}")


@router.get("/document-drafts")
async def get_document_drafts(
        current_user: User = Depends(get_current_active_user)
):
    """Get document drafts from the drafts directory"""
    try:
        # Define drafts directory
        drafts_dir = r"C:\Taitemenetlus\uksikdokumendid\mustandid"
        logger.info(f"Getting document drafts from: {drafts_dir}")

        # Create directory if it doesn't exist
        if not os.path.exists(drafts_dir):
            os.makedirs(drafts_dir, exist_ok=True)
            logger.info(f"Created drafts directory: {drafts_dir}")

        # Get list of draft files
        drafts = []

        try:
            for item in os.listdir(drafts_dir):
                item_path = os.path.join(drafts_dir, item)

                # Skip directories, we only want files
                if os.path.isdir(item_path):
                    continue

                # Get file stats
                stats = os.stat(item_path)

                # Get file extension
                _, extension = os.path.splitext(item)

                # Format the modified time
                modified_time = datetime.fromtimestamp(stats.st_mtime).strftime("%d.%m.%Y %H:%M")

                # Add draft info to the list
                drafts.append({
                    "name": item,
                    "path": item_path,
                    "size": stats.st_size,
                    "formatted_size": format_file_size(stats.st_size),
                    "modified": modified_time,
                    "extension": extension.lower()
                })

            # Sort drafts by name
            drafts.sort(key=lambda x: x["name"].lower())

            return {
                "success": True,
                "drafts_dir": drafts_dir,
                "drafts": drafts
            }

        except Exception as e:
            logger.error(f"Error listing drafts directory: {str(e)}")
            return {
                "success": False,
                "drafts_dir": drafts_dir,
                "error": str(e),
                "drafts": []
            }

    except Exception as e:
        logger.exception(f"Error in get_document_drafts: {str(e)}")
        raise HTTPException(status_code=500, detail=f"Error getting document drafts: {str(e)}")


@router.post("/convert-to-pdf")
async def convert_to_pdf(
        source_path: str = Form(...),
        current_user: User = Depends(get_current_active_user)
):
    """Convert a Word document to PDF"""
    try:
        logger.info(f"Converting document to PDF: {source_path}")

        # Verify the file exists
        if not os.path.exists(source_path):
            return {
                "success": False,
                "message": f"Faili ei leitud: {source_path}"
            }

        # Check if it's a Word document by extension
        _, ext = os.path.splitext(source_path)
        if ext.lower() not in ['.doc', '.docx', '.rtf']:
            return {
                "success": False,
                "message": f"Fail ei ole Word dokument: {source_path}"
            }

        # Create PDF path (same name, different extension)
        pdf_path = os.path.splitext(source_path)[0] + '.pdf'

        # Try to convert the document using LibreOffice (if available)
        try:
            # Check if LibreOffice is available
            libreoffice_paths = [
                r"C:\Program Files\LibreOffice\program\soffice.exe",
                r"C:\Program Files (x86)\LibreOffice\program\soffice.exe",
                "soffice"  # for Linux/macOS
            ]

            libreoffice_path = None
            for path in libreoffice_paths:
                if os.path.exists(path) or path == "soffice":
                    libreoffice_path = path
                    break

            if libreoffice_path:
                # Convert using LibreOffice
                logger.info(f"Converting with LibreOffice: {libreoffice_path}")
                process = subprocess.Popen([
                    libreoffice_path,
                    '--headless',
                    '--convert-to', 'pdf',
                    '--outdir', os.path.dirname(source_path),
                    source_path
                ], stdout=subprocess.PIPE, stderr=subprocess.PIPE)

                stdout, stderr = process.communicate()
                logger.info(f"LibreOffice output: {stdout.decode('utf-8', errors='ignore')}")
                if stderr:
                    logger.warning(f"LibreOffice errors: {stderr.decode('utf-8', errors='ignore')}")

                # Check if the PDF was created
                if os.path.exists(pdf_path):
                    return {
                        "success": True,
                        "message": "PDF dokument edukalt loodud",
                        "pdf_path": pdf_path
                    }
                else:
                    raise Exception("PDF file was not created by LibreOffice")
            else:
                raise Exception("LibreOffice not found")

        except Exception as e:
            logger.error(f"Error using LibreOffice for conversion: {str(e)}")

            # Fallback: Try to use Word via COM automation (Windows only)
            try:
                import win32com.client

                logger.info("Attempting conversion using Word COM automation")
                word = win32com.client.Dispatch("Word.Application")
                word.Visible = False

                doc = word.Documents.Open(source_path)
                doc.SaveAs(pdf_path, FileFormat=17)  # 17 is the PDF format code
                doc.Close()
                word.Quit()

                return {
                    "success": True,
                    "message": "PDF dokument edukalt loodud",
                    "pdf_path": pdf_path
                }

            except Exception as com_error:
                logger.error(f"Error using Word COM automation: {str(com_error)}")
                return {
                    "success": False,
                    "message": f"PDF konverteerimine ebaõnnestus: {str(e)}. Word COM viga: {str(com_error)}"
                }

    except Exception as e:
        logger.exception(f"Error converting to PDF: {str(e)}")
        return {
            "success": False,
            "message": f"Viga PDF konverteerimisel: {str(e)}"
        }


@router.post("/open-for-editing")
async def open_for_editing(
        file_path: str = Form(...),
        current_user: User = Depends(get_current_active_user)
):
    """Open a document for editing in its default application"""
    try:
        logger.info(f"Opening file for editing: {file_path}")

        # Verify the file exists
        if not os.path.exists(file_path):
            return {
                "success": False,
                "message": f"Faili ei leitud: {file_path}"
            }

        # Open the file with the default application
        system = platform.system()

        if system == "Windows":
            os.startfile(file_path)
        elif system == "Darwin":  # macOS
            subprocess.Popen(["open", file_path])
        elif system == "Linux":
            subprocess.Popen(["xdg-open", file_path])
        else:
            return {
                "success": False,
                "message": f"Tundmatu operatsioonisüsteem: {system}"
            }

        return {
            "success": True,
            "message": "Fail avatud redigeerimiseks"
        }

    except Exception as e:
        logger.exception(f"Error opening file for editing: {str(e)}")
        return {
            "success": False,
            "message": f"Viga faili avamisel: {str(e)}"
        }


@router.post("/generate-document")
async def generate_document(
        template_path: str = Form(...),
        row_data_json: str = Form(...),
        current_user: User = Depends(get_current_active_user)
):
    """Generate a document from a template, replacing placeholders with row data values"""
    try:
        logger.info(f"Generating document from template: {template_path}")

        # Parse row data
        try:
            row_data = json.loads(row_data_json)
            logger.info(f"Row data loaded successfully with {len(row_data)} fields")
        except json.JSONDecodeError as e:
            logger.error(f"Error parsing row data JSON: {str(e)}")
            return {
                "success": False,
                "message": f"Viga andmete töötlemisel: {str(e)}"
            }

        # Verify the template file exists
        if not os.path.exists(template_path):
            return {
                "success": False,
                "message": f"Malli faili ei leitud: {template_path}"
            }

        # Get template file extension
        _, ext = os.path.splitext(template_path)
        ext = ext.lower()

        # Determine the target file name for the generated document
        # Try different variations of võlgnik field if available
        volgnik_value = None
        volgnik_keys = ['võlgnik', 'volgnik', 'VÕLGNIK', 'Võlgnik']
        for key in volgnik_keys:
            if key in row_data and row_data[key]:
                volgnik_value = row_data[key]
                break

        # Create a file name based on template name and võlgnik or timestamp
        template_name = os.path.basename(template_path)
        base_name, _ = os.path.splitext(template_name)

        timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
        if volgnik_value:
            # Sanitize the võlgnik value for use in a filename
            volgnik_value = re.sub(r'[\\/:*?"<>|]', '_', str(volgnik_value))
            # Truncate if too long
            if len(volgnik_value) > 30:
                volgnik_value = volgnik_value[:30]
            file_name = f"{base_name}_{volgnik_value}_{timestamp}{ext}"
        else:
            file_name = f"{base_name}_{timestamp}{ext}"

        # Define the drafts directory
        drafts_dir = r"C:\Taitemenetlus\uksikdokumendid\mustandid"

        # Create the directory if it doesn't exist
        if not os.path.exists(drafts_dir):
            os.makedirs(drafts_dir, exist_ok=True)
            logger.info(f"Created drafts directory: {drafts_dir}")

        # Define the output file path
        output_path = os.path.join(drafts_dir, file_name)

        # Process the document based on its type
        success = False

        if ext in ['.docx']:
            # Handle DOCX files using python-docx
            success = await process_docx_template(template_path, output_path, row_data)
        elif ext in ['.doc', '.rtf']:
            # Handle DOC/RTF files using COM automation
            success = await process_doc_template(template_path, output_path, row_data)
        elif ext in ['.txt', '.html', '.xml']:
            # Handle text-based files
            success = await process_text_template(template_path, output_path, row_data)
        else:
            return {
                "success": False,
                "message": f"Ebatoetatud failivorming: {ext}"
            }

        if success:
            logger.info(f"Document generated successfully: {output_path}")
            return {
                "success": True,
                "message": "Dokument edukalt loodud",
                "file_path": output_path,
                "file_name": file_name
            }
        else:
            return {
                "success": False,
                "message": "Dokumendi loomine ebaõnnestus. Kontrollige logisid."
            }

    except Exception as e:
        logger.exception(f"Error generating document: {str(e)}")
        return {
            "success": False,
            "message": f"Viga dokumendi loomisel: {str(e)}"
        }


async def process_docx_template(template_path, output_path, row_data):
    """Process a DOCX template, replacing placeholders with row data values"""
    try:
        # Try to import docx
        try:
            import docx
        except ImportError:
            logger.error("python-docx library not installed. Falling back to COM automation.")
            return await process_doc_template(template_path, output_path, row_data)

        # Load the document
        doc = docx.Document(template_path)

        # Track if any replacements were made
        replacements_made = False

        # Replace placeholders in paragraphs
        for paragraph in doc.paragraphs:
            if replace_text_in_paragraph(paragraph, row_data):
                replacements_made = True

        # Replace placeholders in tables
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    for paragraph in cell.paragraphs:
                        if replace_text_in_paragraph(paragraph, row_data):
                            replacements_made = True

        # Replace placeholders in headers and footers
        for section in doc.sections:
            # Process header
            for paragraph in section.header.paragraphs:
                if replace_text_in_paragraph(paragraph, row_data):
                    replacements_made = True

            # Process footer
            for paragraph in section.footer.paragraphs:
                if replace_text_in_paragraph(paragraph, row_data):
                    replacements_made = True

        # Save the document
        doc.save(output_path)

        logger.info(f"DOCX document processed with {'replacements' if replacements_made else 'no replacements'}")
        return True

    except Exception as e:
        logger.exception(f"Error processing DOCX template: {str(e)}")
        # Try the COM automation as a fallback
        logger.info("Attempting to fall back to COM automation...")
        return await process_doc_template(template_path, output_path, row_data)


def replace_text_in_paragraph(paragraph, row_data):
    """Replace all placeholders in a paragraph with values from row_data. Returns True if replacements were made."""
    # Find all placeholders like <column_name>
    placeholder_pattern = r'<([^>]+)>'

    # Get the paragraph text
    text = paragraph.text

    # Find all matches
    matches = re.findall(placeholder_pattern, text)

    # If no matches, return early
    if not matches:
        return False

    # Create a new text with replacements
    new_text = text
    replacements_made = False

    # Replace each match with the corresponding value from row_data
    for match in matches:
        placeholder = f"<{match}>"

        # Check if the column exists in row_data (case-insensitive search)
        value = None
        match_lower = match.lower()
        for key, val in row_data.items():
            if key.lower() == match_lower and val is not None:
                value = str(val)
                break

        # If value found, replace it
        if value is not None:
            new_text = new_text.replace(placeholder, value)
            replacements_made = True

    # Set the new text back to the paragraph if changes were made
    if replacements_made:
        paragraph.text = new_text

    return replacements_made


async def process_doc_template(template_path, output_path, row_data):
    """Process a DOC/RTF template using COM automation"""
    try:
        # First, copy the template to the output path
        import shutil
        shutil.copy2(template_path, output_path)

        # Try to use Word COM automation
        try:
            import win32com.client
        except ImportError:
            logger.error("pywin32 not installed. Cannot process DOC/RTF files.")
            return False

        logger.info("Using COM automation to process document")
        word = None
        doc = None

        try:
            word = win32com.client.Dispatch("Word.Application")
            word.Visible = False

            # Open the copied document
            doc = word.Documents.Open(output_path)

            # Find and replace placeholders
            replacements_count = 0
            for key, value in row_data.items():
                placeholder = f"<{key}>"
                if value is not None:
                    # Replace placeholder with value
                    find_obj = word.Selection.Find
                    find_obj.ClearFormatting()
                    find_obj.Replacement.ClearFormatting()
                    find_obj.Text = placeholder
                    find_obj.Replacement.Text = str(value)
                    find_obj.Forward = True
                    find_obj.Wrap = 1  # wdFindContinue
                    find_obj.Format = False
                    find_obj.MatchCase = False
                    find_obj.MatchWholeWord = False
                    find_obj.MatchWildcards = False
                    find_obj.MatchSoundsLike = False
                    find_obj.MatchAllWordForms = False

                    # Execute the replacement
                    replacements_in_this_run = 0
                    while find_obj.Execute(FindText=placeholder,
                                           ReplaceWith=str(value),
                                           Replace=1):  # wdReplaceOne
                        replacements_in_this_run += 1
                        replacements_count += 1

            logger.info(f"COM automation: {replacements_count} replacements made")

            # Save and close
            doc.Save()
            return True

        finally:
            # Clean up - properly close Word to avoid orphaned processes
            if doc:
                try:
                    doc.Close(SaveChanges=True)
                except Exception as e:
                    logger.error(f"Error closing document: {str(e)}")

            if word:
                try:
                    word.Quit()
                except Exception as e:
                    logger.error(f"Error quitting Word: {str(e)}")

    except Exception as e:
        logger.exception(f"Error processing DOC template: {str(e)}")
        return False


async def process_text_template(template_path, output_path, row_data):
    """Process a text-based template file"""
    try:
        # Read the template
        with open(template_path, 'r', encoding='utf-8', errors='ignore') as f:
            content = f.read()

        # Replace placeholders
        placeholder_pattern = r'<([^>]+)>'
        matches = re.findall(placeholder_pattern, content)

        replacements_made = False
        for match in matches:
            placeholder = f"<{match}>"

            # Check if the column exists in row_data (case-insensitive search)
            value = None
            match_lower = match.lower()
            for key, val in row_data.items():
                if key.lower() == match_lower and val is not None:
                    value = str(val)
                    break

            # If value found, replace it
            if value is not None:
                content = content.replace(placeholder, value)
                replacements_made = True

        # Write the output
        with open(output_path, 'w', encoding='utf-8') as f:
            f.write(content)

        logger.info(f"Text document processed with {'replacements' if replacements_made else 'no replacements'}")
        return True

    except Exception as e:
        logger.exception(f"Error processing text template: {str(e)}")
        return False


@router.get("/browse-koondaja-folder")
async def browse_koondaja_folder(path: str = ""):
    """Browse the Koondaja folder structure with multiple CSV folders"""
    try:
        base_path = r"C:\TAITEMENETLUS\ÜLESANDED\Tööriistad\ROCKI"

        # Define the allowed folder structure for Koondaja
        koondaja_folders = ["CSV", "Konto vv", "MTA", "Pension", "töötukassa"]

        if path:
            # Handle path normalization
            path = path.replace('/', os.sep).replace('\\', os.sep)
            full_path = os.path.join(base_path, path)
        else:
            # Return the main Koondaja folders
            full_path = base_path

        logger.info(f"Browsing Koondaja folder: {full_path}")

        if not os.path.exists(full_path):
            logger.error(f"Koondaja path does not exist: {full_path}")
            return {
                "success": False,
                "error": f"Path does not exist: {full_path}",
                "items": [],
                "current_path": path
            }

        items = []

        if not path:
            # Return only the Koondaja-specific folders at root level
            for folder_name in koondaja_folders:
                folder_path = os.path.join(base_path, folder_name)
                if os.path.exists(folder_path) and os.path.isdir(folder_path):
                    try:
                        stat_info = os.stat(folder_path)
                        items.append({
                            "name": folder_name,
                            "type": "folder",
                            "path": folder_name,
                            "size": None,
                            "modified": datetime.fromtimestamp(stat_info.st_mtime).strftime("%Y-%m-%d %H:%M:%S")
                        })
                    except Exception as e:
                        logger.warning(f"Error getting folder info for {folder_name}: {str(e)}")
                        continue
        else:
            # Browse within a specific folder
            try:
                for item in os.listdir(full_path):
                    item_path = os.path.join(full_path, item)

                    try:
                        stat_info = os.stat(item_path)
                        is_dir = os.path.isdir(item_path)

                        # For Koondaja, we're primarily interested in CSV files
                        if not is_dir and not item.lower().endswith('.csv'):
                            continue

                        relative_path = os.path.join(path, item) if path else item

                        items.append({
                            "name": item,
                            "type": "folder" if is_dir else "file",
                            "path": relative_path,
                            "size": stat_info.st_size if not is_dir else None,
                            "modified": datetime.fromtimestamp(stat_info.st_mtime).strftime("%Y-%m-%d %H:%M:%S")
                        })
                    except Exception as e:
                        logger.warning(f"Error processing item {item}: {str(e)}")
                        continue

            except PermissionError:
                logger.error(f"Permission denied accessing: {full_path}")
                return {
                    "success": False,
                    "error": f"Permission denied accessing: {full_path}",
                    "items": [],
                    "current_path": path
                }

        # Sort items: folders first, then files, alphabetically
        items.sort(key=lambda x: (x["type"] == "file", x["name"].lower()))

        logger.info(f"Found {len(items)} items in Koondaja folder: {full_path}")

        return {
            "success": True,
            "items": items,
            "current_path": path,
            "base_path": base_path
        }

    except Exception as e:
        logger.exception(f"Error browsing Koondaja folder: {str(e)}")
        return {
            "success": False,
            "error": f"Error browsing folder: {str(e)}",
            "items": [],
            "current_path": path
        }


@router.post("/import-koondaja-csv")
async def import_koondaja_csv(request: Request, db: AsyncSession = Depends(get_db)):
    """Import CSV file for Koondaja data with folder-specific processing and database lookups - ENHANCED VERSION"""
    try:
        # Get the request body
        try:
            body = await request.json()
            logger.info(f"Received Koondaja import request: {body}")
        except Exception as e:
            logger.error(f"Error parsing Koondaja request body: {str(e)}")
            raise HTTPException(status_code=400, detail="Invalid JSON in request body")

        file_path = body.get('file_path')
        folder_type = body.get('folder_type', 'unknown')

        if not file_path:
            logger.error("No file_path provided in Koondaja request")
            raise HTTPException(status_code=400, detail="File path is required")

        # FIXED: Construct the full absolute path
        base_path = r"C:\TAITEMENETLUS\ÜLESANDED\Tööriistad\ROCKI"

        # Handle path normalization and construct full path
        file_path = file_path.replace('/', os.sep).replace('\\', os.sep)
        full_file_path = os.path.join(base_path, file_path)

        logger.info(f"Attempting to import Koondaja CSV file: {full_file_path} (folder_type: {folder_type})")

        if not os.path.exists(full_file_path):
            logger.error(f"Koondaja file not found: {full_file_path}")
            raise HTTPException(status_code=404, detail=f"File not found: {full_file_path}")

        if not os.path.isfile(full_file_path):
            logger.error(f"Koondaja path is not a file: {full_file_path}")
            raise HTTPException(status_code=400, detail=f"Path is not a file: {full_file_path}")

        data = []

        # Try different encodings to handle Estonian characters
        encodings = ['utf-8-sig', 'utf-8', 'windows-1252', 'iso-8859-15', 'cp1257']
        content = None
        used_encoding = None

        for encoding in encodings:
            try:
                logger.debug(f"Trying encoding for Koondaja: {encoding}")
                with open(full_file_path, 'r', encoding=encoding) as csvfile:
                    content = csvfile.read()
                    used_encoding = encoding
                    logger.info(f"Successfully read Koondaja file with encoding: {encoding}")
                    break
            except UnicodeDecodeError as e:
                logger.debug(f"Failed to read Koondaja file with encoding {encoding}: {str(e)}")
                continue
            except Exception as e:
                logger.error(f"Unexpected error reading Koondaja file with encoding {encoding}: {str(e)}")
                continue

        if content is None:
            logger.error("Could not read Koondaja file with any encoding")
            raise HTTPException(status_code=400, detail="Could not decode file with any supported encoding")

        # Parse CSV content
        try:
            # Try to detect delimiter
            csv_sniffer = csv.Sniffer()
            delimiter = csv_sniffer.sniff(content[:1024]).delimiter
            logger.info(f"Detected delimiter for Koondaja: '{delimiter}'")
        except Exception as e:
            logger.warning(f"Could not detect delimiter for Koondaja, using default ';': {str(e)}")
            delimiter = ';'

        try:
            csv_reader = csv.reader(StringIO(content), delimiter=delimiter)
            rows = list(csv_reader)

            logger.info(f"Koondaja CSV parsing successful. Total rows: {len(rows)}")

            if not rows:
                logger.warning("Koondaja CSV file is empty")
                return {
                    "success": True,
                    "message": "File is empty",
                    "data": [],
                    "folder_type": folder_type,
                    "encoding_used": used_encoding,
                    "total_rows_processed": 0,
                    "valid_rows": 0,
                    "lookup_data": {"viitenumber_lookup": {}, "isikukood_lookup": {}}
                }

            # Process rows based on folder type
            processed_rows = 0
            row_count = len(rows)

            for row_num, row in enumerate(rows, 1):
                try:
                    # Skip empty rows
                    if not any(cell.strip() for cell in row if cell):
                        continue

                    # Process based on folder type
                    if folder_type == 'konto_vv':
                        processed_row = process_konto_vv_row(row, row_num)
                    elif folder_type == 'csv':
                        processed_row = process_csv_folder_row(row, row_num)
                    elif folder_type == 'mta':
                        processed_row = process_mta_row(row, row_num)
                    elif folder_type == 'pension':
                        processed_row = process_pension_row(row, row_num)
                    elif folder_type == 'tootukassa':
                        processed_row = process_tootukassa_row(row, row_num)
                    else:
                        # Default processing - just return the raw row as array
                        processed_row = row

                    if processed_row is not None:
                        data.append(processed_row)
                        processed_rows += 1

                except Exception as e:
                    logger.warning(f"Error processing Koondaja row {row_num}: {str(e)}")
                    continue

            logger.info(
                f"Successfully processed {processed_rows} out of {row_count} Koondaja rows from {os.path.basename(full_file_path)}")

            # ENHANCED: Perform database lookups for optimal performance
            lookup_data = {}
            if folder_type == 'konto_vv' and data:
                try:
                    logger.info("Performing batch database lookups for Konto vv data...")
                    lookup_data = await perform_koondaja_database_lookups(db, data)
                    logger.info(
                        f"Database lookups completed: {len(lookup_data.get('viitenumber_lookup', {}))} viitenumber matches, {len(lookup_data.get('isikukood_lookup', {}))} isikukood matches")
                except Exception as e:
                    logger.error(f"Error performing database lookups: {str(e)}")
                    lookup_data = {"viitenumber_lookup": {}, "isikukood_lookup": {}}
            else:
                lookup_data = {"viitenumber_lookup": {}, "isikukood_lookup": {}}

            return {
                "success": True,
                "message": f"Successfully imported {len(data)} rows from {os.path.basename(full_file_path)}",
                "data": data,
                "folder_type": folder_type,
                "encoding_used": used_encoding,
                "total_rows_processed": row_count,
                "valid_rows": len(data),
                "lookup_data": lookup_data
            }

        except Exception as e:
            logger.error(f"Error parsing Koondaja CSV content: {str(e)}")
            raise HTTPException(status_code=400, detail=f"Error parsing CSV file: {str(e)}")

    except HTTPException:
        raise
    except Exception as e:
        logger.exception(f"Unexpected error importing Koondaja CSV: {str(e)}")
        raise HTTPException(status_code=500, detail=f"Unexpected error importing CSV: {str(e)}")


def process_konto_vv_row(row: list, row_num: int) -> list:
    """Process a row from Konto vv folder according to specific mapping"""
    try:
        # Ensure row has enough columns by padding with empty strings
        while len(row) < 20:
            row.append('')

        # Clean up the row data - strip whitespace and handle None values
        cleaned_row = []
        for i, cell in enumerate(row):
            if cell is None:
                cleaned_row.append('')
            else:
                cleaned_row.append(str(cell).strip())

        logger.debug(f"Processing Konto vv row {row_num}: columns 10-13 = {cleaned_row[10:14]}")

        # Return the cleaned row as array - transformation will happen in frontend
        return cleaned_row

    except Exception as e:
        logger.warning(f"Error processing Konto vv row {row_num}: {str(e)}")
        return None


def process_csv_folder_row(row: list, row_num: int) -> list:
    """Process a row from CSV folder - placeholder for future implementation"""
    try:
        # For now, return raw row - will be implemented later
        return row
    except Exception as e:
        logger.warning(f"Error processing CSV folder row {row_num}: {str(e)}")
        return None


def process_mta_row(row: list, row_num: int) -> list:
    """Process a row from MTA folder - placeholder for future implementation"""
    try:
        # For now, return raw row - will be implemented later
        return row
    except Exception as e:
        logger.warning(f"Error processing MTA row {row_num}: {str(e)}")
        return None


def process_pension_row(row: list, row_num: int) -> list:
    """Process a row from Pension folder - placeholder for future implementation"""
    try:
        # For now, return raw row - will be implemented later
        return row
    except Exception as e:
        logger.warning(f"Error processing Pension row {row_num}: {str(e)}")
        return None


def process_tootukassa_row(row: list, row_num: int) -> list:
    """Process a row from töötukassa folder - placeholder for future implementation"""
    try:
        # For now, return raw row - will be implemented later
        return row
    except Exception as e:
        logger.warning(f"Error processing töötukassa row {row_num}: {str(e)}")
        return None


async def perform_koondaja_database_lookups(db: AsyncSession, csv_data: list) -> dict:
    """
    Perform batch database lookups for Koondaja data with optimal performance
    Returns lookup maps for viitenumber and võlgniku_isikukood
    """
    try:
        logger.info(f"Starting batch database lookups for {len(csv_data)} CSV rows")

        # Extract unique values for batch lookup
        viitenumbrid = set()
        isikukoodid = set()

        for row in csv_data:
            if isinstance(row, list) and len(row) > 14:
                # Viitenumber from 10th item (index 9)
                viitenumber = str(row[9] if len(row) > 9 else '').strip()
                if viitenumber:
                    viitenumbrid.add(viitenumber)

                # Isikukood from 15th item (index 14)
                isikukood = str(row[14] if len(row) > 14 else '').strip()
                if isikukood:
                    isikukoodid.add(isikukood)

        logger.info(
            f"Found {len(viitenumbrid)} unique viitenumbrid and {len(isikukoodid)} unique isikukoodid for lookup")

        # Initialize lookup maps
        viitenumber_lookup = {}
        isikukood_lookup = {}

        # Get table information first
        try:
            # Get table schema information
            schema_query = text("""
                SELECT name FROM sqlite_master 
                WHERE type='table' AND name LIKE '%table%' OR name LIKE '%main%'
                ORDER BY name
            """)
            schema_result = await db.execute(schema_query)
            tables = [row[0] for row in schema_result.fetchall()]
            logger.info(f"Available tables: {tables}")

            # Try to find the main table name
            main_table_name = None
            for table in tables:
                if 'main' in table.lower() or len(tables) == 1:
                    main_table_name = table
                    break

            if not main_table_name and tables:
                main_table_name = tables[0]  # Use first table as fallback

            logger.info(f"Using table: {main_table_name}")

        except Exception as e:
            logger.warning(f"Could not determine table name, using 'main_table': {str(e)}")
            main_table_name = "main_table"

        # Batch lookup for viitenumber -> toimiku_nr
        if viitenumbrid and main_table_name:
            try:
                viitenumber_list = list(viitenumbrid)

                # Create parameterized query
                placeholders = ','.join([f':vn{i}' for i in range(len(viitenumber_list))])
                params = {f'vn{i}': vn for i, vn in enumerate(viitenumber_list)}

                # Try different possible column combinations
                viitenumber_queries = [
                    f'SELECT "Viitenumber", "Toimiku nr lõplik" FROM "{main_table_name}" WHERE "Viitenumber" IN ({placeholders})',
                    f'SELECT Viitenumber, "Toimiku nr lõplik" FROM "{main_table_name}" WHERE Viitenumber IN ({placeholders})',
                    f'SELECT "Viitenumber", toimiku_nr FROM "{main_table_name}" WHERE "Viitenumber" IN ({placeholders})',
                    f'SELECT viitenumber, toimiku_nr FROM "{main_table_name}" WHERE viitenumber IN ({placeholders})'
                ]

                for query in viitenumber_queries:
                    try:
                        logger.debug(f"Trying viitenumber query: {query}")
                        result = await db.execute(text(query), params)
                        rows = result.fetchall()

                        for row in rows:
                            key = str(row[0]).strip() if row[0] else ''
                            value = str(row[1]).strip() if row[1] else ''
                            if key:
                                viitenumber_lookup[key] = value

                        logger.info(f"Viitenumber lookup successful: {len(viitenumber_lookup)} matches found")
                        break

                    except Exception as e:
                        logger.debug(f"Viitenumber query failed: {str(e)}")
                        continue

            except Exception as e:
                logger.warning(f"Error in viitenumber batch lookup: {str(e)}")

        # Batch lookup for isikukood -> toimiku_nr
        if isikukoodid and main_table_name:
            try:
                isikukood_list = list(isikukoodid)

                # Create parameterized query
                placeholders = ','.join([f':ik{i}' for i in range(len(isikukood_list))])
                params = {f'ik{i}': ik for i, ik in enumerate(isikukood_list)}

                # Try different possible column combinations
                isikukood_queries = [
                    f'SELECT "Võlgniku reg isikukood", "Toimiku nr lõplik" FROM "{main_table_name}" WHERE "Võlgniku reg isikukood" IN ({placeholders})',
                    f'SELECT "Isiku- või registrikood", "Toimiku nr lõplik" FROM "{main_table_name}" WHERE "Isiku- või registrikood" IN ({placeholders})',
                    f'SELECT "Võlgniku reg isikukood", toimiku_nr FROM "{main_table_name}" WHERE "Võlgniku reg isikukood" IN ({placeholders})',
                    f'SELECT "Isiku- või registrikood", toimiku_nr FROM "{main_table_name}" WHERE "Isiku- või registrikood" IN ({placeholders})',
                    f'SELECT volgniku_isikukood, "Toimiku nr lõplik" FROM "{main_table_name}" WHERE volgniku_isikukood IN ({placeholders})',
                    f'SELECT isikukood, "Toimiku nr lõplik" FROM "{main_table_name}" WHERE isikukood IN ({placeholders})'
                ]

                for query in isikukood_queries:
                    try:
                        logger.debug(f"Trying isikukood query: {query}")
                        result = await db.execute(text(query), params)
                        rows = result.fetchall()

                        for row in rows:
                            key = str(row[0]).strip() if row[0] else ''
                            value = str(row[1]).strip() if row[1] else ''
                            if key:
                                isikukood_lookup[key] = value

                        logger.info(f"Isikukood lookup successful: {len(isikukood_lookup)} matches found")
                        break

                    except Exception as e:
                        logger.debug(f"Isikukood query failed: {str(e)}")
                        continue

            except Exception as e:
                logger.warning(f"Error in isikukood batch lookup: {str(e)}")

        logger.info(
            f"Database lookups completed: {len(viitenumber_lookup)} viitenumber matches, {len(isikukood_lookup)} isikukood matches")

        return {
            "viitenumber_lookup": viitenumber_lookup,
            "isikukood_lookup": isikukood_lookup
        }

    except Exception as e:
        logger.error(f"Error performing database lookups: {str(e)}")
        return {
            "viitenumber_lookup": {},
            "isikukood_lookup": {}
        }
